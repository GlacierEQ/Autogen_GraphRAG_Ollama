"""
Microsoft Word document processor for GraphRAG integration.
Allows ingestion of .docx files into the knowledge graph.
"""
import os
import shutil
import tempfile
from typing import Dict, List, Optional, Any
import docx
from docx import Document
import re
import logging

logger = logging.getLogger(__name__)

class WordDocumentProcessor:
    """Process Word documents for GraphRAG indexing."""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """Initialize the Word document processor with optional configuration."""
        self.config = config or {}
        self.chunk_size = int(self.config.get("chunkSize", 300))
        self.chunk_overlap = int(self.config.get("chunkOverlap", 100))
        self.include_images = bool(self.config.get("includeImages", False))
    
    def extract_text(self, file_path: str) -> str:
        """Extract text content from a Word document."""
        try:
            doc = Document(file_path)
            full_text = []
            images = []  # To store image paths if needed
            
            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    full_text.append(para.text)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    full_text.append(" | ".join(row_text))
            
            # Process images if required
            if self.include_images:
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        img_path = os.path.join(tempfile.gettempdir(), rel.target_ref.split("/")[-1])
                        with open(img_path, "wb") as img_file:
                            img_file.write(rel.target_part.blob)
                        images.append(img_path)
            
            return "\n\n".join(full_text), images
        except Exception as e:
            logger.error(f"Error extracting text from Word document: {str(e)}")
            return f"Error: Could not extract text from document ({str(e)})", []

    def extract_metadata(self, file_path: str) -> Dict[str, str]:
        """Extract metadata from a Word document."""
        try:
            doc = Document(file_path)
            metadata = {}
            
            # Core properties
            core_props = doc.core_properties
            
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.created:
                metadata["created"] = str(core_props.created)
            if core_props.modified:
                metadata["modified"] = str(core_props.modified)
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
            
            # Document statistics
            metadata["word_count"] = str(len(re.findall(r'\b\w+\b', self.extract_text(file_path))))
            metadata["file_size"] = str(os.path.getsize(file_path))
            
            return metadata
        except Exception as e:
            logger.error(f"Error extracting metadata from Word document: {str(e)}")
            return {"error": f"Could not extract metadata ({str(e)})"}
    
    def convert_to_markdown(self, file_path: str, output_dir: str) -> str:
        """Convert Word document to Markdown format for GraphRAG indexing."""
        filename = os.path.basename(file_path)
        name, _ = os.path.splitext(filename)
        output_path = os.path.join(output_dir, f"{name}.md")
        
        try:
            doc = Document(file_path)
            images = []  # To store image paths
            
            with open(output_path, "w", encoding="utf-8") as md_file:
                # Add metadata as YAML frontmatter
                metadata = self.extract_metadata(file_path)
                md_file.write("---\n")
                for key, value in metadata.items():
                    if value and value.strip():
                        md_file.write(f"{key}: {value}\n")
                md_file.write("---\n\n")
                
                # Process headings and paragraphs
                for para in doc.paragraphs:
                    if not para.text.strip():
                        continue
                        
                    # Handle headings
                    if para.style.name.startswith('Heading'):
                        level = int(para.style.name.replace('Heading', ''))
                        md_file.write(f"{'#' * level} {para.text}\n\n")
                    else:
                        # Handle formatting
                        text = para.text
                        
                        # Bold
                        for run in para.runs:
                            if run.bold:
                                bold_text = f"**{run.text}**"
                                text = text.replace(run.text, bold_text)
                        
                        # Italic
                        for run in para.runs:
                            if run.italic:
                                italic_text = f"*{run.text}*"
                                text = text.replace(run.text, italic_text)
                        
                        md_file.write(f"{text}\n\n")
                
                # Handle tables
                for table in doc.tables:
                    # Create table header
                    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                    md_file.write("| " + " | ".join(header_cells) + " |\n")
                    md_file.write("| " + " | ".join(["---"] * len(header_cells)) + " |\n")
                    
                    # Create table rows
                    for row in table.rows[1:]:
                        cell_texts = [cell.text.strip() for cell in row.cells]
                        md_file.write("| " + " | ".join(cell_texts) + " |\n")
                    
                    md_file.write("\n")
                
                # Handle images
                for img_path in images:
                    md_file.write(f"![Image]({img_path})\n\n")
            
            return output_path
        except Exception as e:
            logger.error(f"Error converting Word document to Markdown: {str(e)}")
            return f"Error: Could not convert document to Markdown ({str(e)})"

    def process_directory(self, input_dir: str, output_dir: str) -> List[str]:
        """Process all Word documents in a directory."""
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        processed_files = []
        
        for root, _, files in os.walk(input_dir):
            for file in files:
                if file.lower().endswith(('.docx', '.doc')):
                    file_path = os.path.join(root, file)
                    result = self.convert_to_markdown(file_path, output_dir)
                    if not result.startswith("Error"):
                        processed_files.append(result)
        
        return processed_files

# Command line interface
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Process Word documents for GraphRAG")
    parser.add_argument("--input", "-i", required=True, help="Input directory or file")
    parser.add_argument("--output", "-o", required=True, help="Output directory for Markdown files")
    args = parser.parse_args()
    
    processor = WordDocumentProcessor()
    
    if os.path.isdir(args.input):
        results = processor.process_directory(args.input, args.output)
        print(f"Processed {len(results)} files")
        for result in results:
            print(f"  - {result}")
    else:
        result = processor.convert_to_markdown(args.input, args.output)
        print(f"Result: {result}")
